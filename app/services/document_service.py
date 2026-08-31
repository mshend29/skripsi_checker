from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
from difflib import SequenceMatcher
from pathlib import Path

import pymupdf
from docx import Document as DocxDocument


BASE_DIR = Path(__file__).resolve().parents[2]
STORAGE_DIR = BASE_DIR / "storage"
PREVIEW_DIR = STORAGE_DIR / "previews"

DOCUMENT_FOLDERS = {
    "proposal": STORAGE_DIR / "proposals",
    "revision": STORAGE_DIR / "revisions",
    "final": STORAGE_DIR / "finals",
}

MAJOR_SECTION_NAMES = {
    "ABSTRAK",
    "ABSTRACT",
    "DAFTAR PUSTAKA",
    "REFERENCES",
    "LAMPIRAN",
    "APPENDIX",
}


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        document = DocxDocument(path)
        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

    if suffix == ".pdf":
        with pymupdf.open(path) as document:
            return "\n".join(page.get_text("text") for page in document)

    raise ValueError("Format dokumen belum didukung. Gunakan DOCX atau PDF.")


def extract_sections(text: str) -> list[str]:
    lines = [_normalize_text(line) for line in text.splitlines()]
    sections: list[str] = []
    seen: set[str] = set()

    for line in lines:
        if not line or len(line) > 180:
            continue
        if not _heading_label(line):
            continue

        key = line.casefold()
        if key in seen:
            continue

        seen.add(key)
        sections.append(line)

        if len(sections) >= 150:
            break

    return sections


def detect_title(text: str) -> str:
    lines = [_normalize_text(line) for line in text.splitlines()]
    lines = [line for line in lines if line]

    candidates: list[str] = []
    ignored = {
        "proposal",
        "proposal skripsi",
        "skripsi",
        "tugas akhir",
        "bab i",
        "pendahuluan",
        "daftar isi",
    }

    for line in lines[:50]:
        normalized = line.casefold()
        if normalized in ignored:
            continue
        if normalized.startswith("bab i"):
            break
        if 15 <= len(line) <= 220:
            candidates.append(line)

    if not candidates:
        return ""

    return max(candidates, key=len)


def resolve_storage_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.is_absolute():
        return path
    return BASE_DIR / path


def store_document(
    thesis_id: int,
    source_path: str | Path,
    version: int,
    kind: str,
) -> tuple[str, str]:
    if kind not in DOCUMENT_FOLDERS:
        raise ValueError("Jenis dokumen tidak dikenali.")

    source = Path(source_path)
    if not source.is_file():
        raise ValueError("File sumber tidak ditemukan.")

    if source.suffix.lower() not in {".docx", ".pdf"}:
        raise ValueError("Dokumen harus berupa file DOCX atau PDF.")

    target_dir = DOCUMENT_FOLDERS[kind] / str(thesis_id)
    target_dir.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r"[^A-Za-z0-9._ -]", "_", source.name)
    target = target_dir / f"v{version}_{safe_name}"

    if target.exists():
        counter = 2
        while target.exists():
            target = target_dir / f"v{version}_{counter}_{safe_name}"
            counter += 1

    shutil.copy2(source, target)

    text = extract_text(target)
    detected_title = detect_title(text)

    relative_path = target.relative_to(BASE_DIR).as_posix()
    return relative_path, detected_title


def import_proposal(
    thesis_id: int,
    source_path: str | Path,
    version: int,
) -> tuple[str, str]:
    return store_document(
        thesis_id=thesis_id,
        source_path=source_path,
        version=version,
        kind="proposal",
    )


def _heading_label(text: str, style_name: str = "") -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return False

    style_lower = style_name.lower()
    if style_lower.startswith("heading"):
        return True

    patterns = [
        re.compile(
            r"^BAB\s+(?:[IVXLCDM]+|\d+)(?:\s*[-.:]?\s*.*)?$",
            re.IGNORECASE,
        ),
        re.compile(r"^\d+(?:\.\d+){1,3}\s+.+$"),
    ]
    if any(pattern.match(normalized) for pattern in patterns):
        return True

    return normalized.upper() in MAJOR_SECTION_NAMES


def _chapter_key(text: str) -> str | None:
    normalized = _normalize_text(text)
    match = re.match(
        r"^(BAB\s+(?:[IVXLCDM]+|\d+))\b",
        normalized,
        re.IGNORECASE,
    )
    if match:
        return match.group(1).upper()
    return None


def _numbering_key(text: str) -> str | None:
    normalized = _normalize_text(text)
    match = re.match(r"^(\d+(?:\.\d+){1,3})\b", normalized)
    return match.group(1) if match else None


def _structure_level(text: str, style_name: str = "") -> int:
    style_match = re.search(r"(\d+)", style_name or "")
    if style_name.lower().startswith("toc") and style_match:
        return max(1, min(4, int(style_match.group(1))))

    if _chapter_key(text):
        return 1

    numbering = _numbering_key(text)
    if numbering:
        return min(4, numbering.count(".") + 1)

    return 1


def _looks_like_toc_entry(raw_text: str, style_name: str = "") -> bool:
    stripped = (raw_text or "").strip()
    if not stripped:
        return False

    if style_name.lower().startswith("toc"):
        return True

    if "\t" in stripped:
        return True

    if re.search(r"\.{3,}\s*(?:\d+|[ivxlcdm]+)\s*$", stripped, re.I):
        return True

    if (
        (_chapter_key(stripped) or _numbering_key(stripped))
        and re.search(r"\s+(?:\d+|[ivxlcdm]+)\s*$", stripped, re.I)
    ):
        return True

    return False


def _clean_toc_entry(raw_text: str) -> str:
    value = (raw_text or "").replace("\xa0", " ").strip()
    value = re.sub(
        r"\s*(?:\.{2,}|\t+)\s*(?:\d+|[ivxlcdm]+)\s*$",
        "",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        r"(?<=\D)\s{2,}(?:\d+|[ivxlcdm]+)\s*$",
        "",
        value,
        flags=re.IGNORECASE,
    )
    return _normalize_text(value)


def _find_toc_bounds(rows: list[dict]) -> tuple[int | None, int | None]:
    toc_start: int | None = None

    for position, row in enumerate(rows):
        if row["text"].casefold() == "daftar isi":
            toc_start = position
            break

    if toc_start is None:
        return None, None

    toc_end = toc_start
    saw_entry = False

    for position in range(toc_start + 1, len(rows)):
        row = rows[position]
        raw_text = row.get("raw_text") or row["text"]
        style_name = row.get("style_name") or ""

        if _looks_like_toc_entry(raw_text, style_name):
            saw_entry = True
            toc_end = position
            continue

        if not saw_entry:
            if position - toc_start <= 3:
                toc_end = position
                continue
            break

        if _chapter_key(row["text"]) or (
            style_name.lower().startswith("heading")
            and not style_name.lower().startswith("toc")
        ):
            break

        if position - toc_end <= 2 and len(row["text"]) < 120:
            toc_end = position
            continue

        break

    return toc_start, toc_end


def _match_structure_target(
    label: str,
    rows: list[dict],
    start_position: int,
) -> int | None:
    chapter_key = _chapter_key(label)
    numbering_key = _numbering_key(label)
    normalized_label = _normalize_text(label).casefold()

    for row in rows[start_position:]:
        text = row["text"]
        text_lower = text.casefold()

        if chapter_key and _chapter_key(text) == chapter_key:
            return int(row["paragraph_index"])

        if numbering_key and _numbering_key(text) == numbering_key:
            return int(row["paragraph_index"])

        if normalized_label == text_lower:
            return int(row["paragraph_index"])

        if len(normalized_label) > 8 and (
            normalized_label in text_lower or text_lower in normalized_label
        ):
            return int(row["paragraph_index"])

    return None


def _build_tree(entries: list[dict]) -> list[dict]:
    roots: list[dict] = []
    stack: list[tuple[int, dict]] = []

    for entry in entries:
        node = {
            "label": entry["label"],
            "paragraph_index": entry.get("paragraph_index"),
            "children": [],
        }
        level = max(1, int(entry.get("level", 1)))

        while stack and stack[-1][0] >= level:
            stack.pop()

        if stack:
            stack[-1][1]["children"].append(node)
        else:
            roots.append(node)

        stack.append((level, node))

    return roots


def extract_review_paragraphs(file_path: str | Path) -> list[dict]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    rows: list[dict] = []

    if suffix == ".docx":
        document = DocxDocument(path)
        for source_index, paragraph in enumerate(document.paragraphs):
            raw_text = paragraph.text.strip()
            text_value = _normalize_text(raw_text)
            if not text_value:
                continue

            style_name = paragraph.style.name if paragraph.style else ""
            rows.append(
                {
                    "paragraph_index": source_index,
                    "display_index": len(rows) + 1,
                    "text": text_value,
                    "raw_text": raw_text,
                    "section": "Halaman Awal / Halaman Judul",
                    "style_name": style_name,
                    "is_heading": _heading_label(text_value, style_name),
                }
            )

    elif suffix == ".pdf":
        with pymupdf.open(path) as document:
            source_index = 0
            for page_number, page in enumerate(document, start=1):
                blocks = page.get_text("blocks")
                for block in blocks:
                    raw_text = str(block[4]).strip()
                    text_value = _normalize_text(raw_text)
                    if not text_value:
                        continue

                    rows.append(
                        {
                            "paragraph_index": source_index,
                            "display_index": len(rows) + 1,
                            "text": text_value,
                            "raw_text": raw_text,
                            "section": "Halaman Awal / Halaman Judul",
                            "style_name": f"PDF halaman {page_number}",
                            "is_heading": _heading_label(text_value),
                        }
                    )
                    source_index += 1
    else:
        raise ValueError(
            "Format dokumen belum didukung. Gunakan DOCX atau PDF."
        )

    toc_start, toc_end = _find_toc_bounds(rows)
    current_section = "Halaman Awal / Halaman Judul"

    for position, row in enumerate(rows):
        if toc_start is not None and toc_end is not None:
            if toc_start <= position <= toc_end:
                row["section"] = "Daftar Isi"
                continue
            if position < toc_start:
                row["section"] = "Halaman Awal / Halaman Judul"
                continue

        if row["is_heading"]:
            current_section = row["text"]

        row["section"] = current_section

    return rows


def extract_review_structure(
    file_path: str | Path,
    paragraphs: list[dict] | None = None,
) -> list[dict]:
    rows = paragraphs or extract_review_paragraphs(file_path)
    if not rows:
        return []

    roots: list[dict] = [
        {
            "label": "Halaman Awal / Halaman Judul",
            "paragraph_index": int(rows[0]["paragraph_index"]),
            "children": [],
        }
    ]

    toc_start, toc_end = _find_toc_bounds(rows)
    entries: list[dict] = []

    if toc_start is not None and toc_end is not None:
        roots.append(
            {
                "label": "Daftar Isi",
                "paragraph_index": int(rows[toc_start]["paragraph_index"]),
                "children": [],
            }
        )

        search_start = min(len(rows), toc_end + 1)

        for row in rows[toc_start + 1 : toc_end + 1]:
            raw_text = row.get("raw_text") or row["text"]
            style_name = row.get("style_name") or ""
            if not _looks_like_toc_entry(raw_text, style_name):
                continue

            label = _clean_toc_entry(raw_text)
            if not label:
                continue

            label_folded = label.casefold()
            if label_folded in {
                "daftar isi",
                "halaman awal",
                "halaman judul",
            }:
                continue

            if not (
                _chapter_key(label)
                or _numbering_key(label)
                or label.upper() in MAJOR_SECTION_NAMES
            ):
                continue

            entries.append(
                {
                    "label": label,
                    "level": _structure_level(label, style_name),
                    "paragraph_index": _match_structure_target(
                        label,
                        rows,
                        search_start,
                    ),
                }
            )

        if entries:
            roots.extend(_build_tree(entries))
            return roots

    fallback_entries: list[dict] = []
    start_position = (toc_end + 1) if toc_end is not None else 0

    for row in rows[start_position:]:
        label = row["text"]
        if label.casefold() == "daftar isi":
            if not any(root["label"] == "Daftar Isi" for root in roots):
                roots.append(
                    {
                        "label": "Daftar Isi",
                        "paragraph_index": int(row["paragraph_index"]),
                        "children": [],
                    }
                )
            continue

        if not (
            _chapter_key(label)
            or _numbering_key(label)
            or label.upper() in MAJOR_SECTION_NAMES
        ):
            continue

        fallback_entries.append(
            {
                "label": label,
                "level": _structure_level(
                    label,
                    row.get("style_name") or "",
                ),
                "paragraph_index": int(row["paragraph_index"]),
            }
        )

    roots.extend(_build_tree(fallback_entries))
    return roots


def export_docx_with_comments(
    source_path: str | Path,
    output_path: str | Path,
    comments: list[dict],
    author: str = "Dosen Pembimbing",
    initials: str = "DP",
) -> int:
    source = Path(source_path)
    target = Path(output_path)

    if source.suffix.lower() != ".docx":
        raise ValueError(
            "Komentar Word hanya dapat ditanamkan ke dokumen DOCX."
        )

    document = DocxDocument(source)
    exported = 0

    for item in comments:
        paragraph_index = item.get("paragraph_index")
        if paragraph_index is None:
            continue

        if not 0 <= int(paragraph_index) < len(document.paragraphs):
            continue

        paragraph = document.paragraphs[int(paragraph_index)]
        runs = list(paragraph.runs)

        if not runs:
            runs = [paragraph.add_run("")]

        category = item.get("category") or "Umum"
        severity = item.get("severity") or "Moderate"
        selected_text = (item.get("selected_text") or "").strip()
        comment_text = (item.get("content") or "").strip()

        parts = [f"[{severity}] [{category}]"]
        if selected_text and selected_text != paragraph.text.strip():
            parts.append(f'Kutipan: "{selected_text}"')
        parts.append(comment_text)

        document.add_comment(
            runs=runs,
            text="\n".join(parts),
            author=author,
            initials=initials,
        )
        exported += 1

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return exported


def _preview_cache_path(source: Path) -> Path:
    stat = source.stat()
    payload = f"{source.resolve()}|{stat.st_size}|{stat.st_mtime_ns}"
    digest = hashlib.sha1(payload.encode("utf-8")).hexdigest()[:14]
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    return PREVIEW_DIR / f"{source.stem}_{digest}.pdf"


def _convert_docx_with_word(source: Path, target: Path) -> None:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    document = None

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        document = word.Documents.Open(
            str(source.resolve()),
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        document.ExportAsFixedFormat(
            OutputFileName=str(target.resolve()),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            Range=0,
            Item=0,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=1,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _convert_docx_with_libreoffice(source: Path, target: Path) -> None:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise RuntimeError("LibreOffice tidak ditemukan.")

    target.parent.mkdir(parents=True, exist_ok=True)

    result = subprocess.run(
        [
            executable,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target.parent),
            str(source.resolve()),
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    generated = target.parent / f"{source.stem}.pdf"
    if result.returncode != 0 or not generated.exists():
        message = (result.stderr or result.stdout or "").strip()
        raise RuntimeError(
            f"LibreOffice gagal membuat preview PDF. {message}"
        )

    if generated.resolve() != target.resolve():
        if target.exists():
            target.unlink()
        generated.replace(target)


def ensure_preview_pdf(source_path: str | Path) -> Path:
    source = resolve_storage_path(source_path)

    if not source.exists():
        raise FileNotFoundError(str(source))

    if source.suffix.lower() == ".pdf":
        return source

    if source.suffix.lower() != ".docx":
        raise ValueError(
            "Preview visual saat ini mendukung DOCX dan PDF."
        )

    target = _preview_cache_path(source)
    if target.exists() and target.stat().st_size > 0:
        return target

    errors: list[str] = []

    if os.name == "nt":
        try:
            _convert_docx_with_word(source, target)
            if target.exists() and target.stat().st_size > 0:
                return target
        except Exception as exc:
            errors.append(f"Microsoft Word: {exc}")

    try:
        _convert_docx_with_libreoffice(source, target)
        if target.exists() and target.stat().st_size > 0:
            return target
    except Exception as exc:
        errors.append(f"LibreOffice: {exc}")

    details = "\n".join(errors) or "Tidak ada converter DOCX tersedia."
    raise RuntimeError(
        "Preview Word dengan layout asli membutuhkan Microsoft Word "
        "atau LibreOffice yang terpasang di komputer.\n\n"
        + details
    )


def render_preview_pages(
    preview_pdf: str | Path,
    dpi: int = 120,
) -> list[dict]:
    path = Path(preview_pdf)
    pages: list[dict] = []
    zoom = dpi / 72.0
    matrix = pymupdf.Matrix(zoom, zoom)

    with pymupdf.open(path) as document:
        for page_index, page in enumerate(document):
            pixmap = page.get_pixmap(
                matrix=matrix,
                alpha=False,
            )
            pages.append(
                {
                    "page_index": page_index,
                    "page_width": float(page.rect.width),
                    "page_height": float(page.rect.height),
                    "image_width": int(pixmap.width),
                    "image_height": int(pixmap.height),
                    "png": pixmap.tobytes("png"),
                }
            )

    return pages


def _pdf_text_blocks(preview_pdf: Path) -> list[dict]:
    result: list[dict] = []

    with pymupdf.open(preview_pdf) as document:
        for page_index, page in enumerate(document):
            words = page.get_text("words")
            page_words = [
                {
                    "x0": float(word[0]),
                    "y0": float(word[1]),
                    "x1": float(word[2]),
                    "y1": float(word[3]),
                    "text": str(word[4]),
                    "block_no": int(word[5]),
                    "line_no": int(word[6]),
                    "word_no": int(word[7]),
                }
                for word in words
            ]

            for block in page.get_text("blocks"):
                text_value = _normalize_text(str(block[4]))
                if not text_value:
                    continue

                bbox = (
                    float(block[0]),
                    float(block[1]),
                    float(block[2]),
                    float(block[3]),
                )
                x0, y0, x1, y1 = bbox

                block_words = [
                    word
                    for word in page_words
                    if (
                        ((word["x0"] + word["x1"]) / 2) >= x0 - 1
                        and ((word["x0"] + word["x1"]) / 2) <= x1 + 1
                        and ((word["y0"] + word["y1"]) / 2) >= y0 - 1
                        and ((word["y0"] + word["y1"]) / 2) <= y1 + 1
                    )
                ]

                result.append(
                    {
                        "page_index": page_index,
                        "bbox": bbox,
                        "text": text_value,
                        "words": block_words,
                    }
                )

    return result


def _union_bbox(blocks: list[dict]) -> tuple[float, float, float, float]:
    return (
        min(block["bbox"][0] for block in blocks),
        min(block["bbox"][1] for block in blocks),
        max(block["bbox"][2] for block in blocks),
        max(block["bbox"][3] for block in blocks),
    )


def _merge_words(blocks: list[dict]) -> list[dict]:
    words: list[dict] = []
    seen: set[tuple] = set()

    for block in blocks:
        for word in block.get("words") or []:
            key = (
                round(float(word["x0"]), 2),
                round(float(word["y0"]), 2),
                round(float(word["x1"]), 2),
                round(float(word["y1"]), 2),
                word["text"],
            )
            if key in seen:
                continue
            seen.add(key)
            words.append(dict(word))

    words.sort(
        key=lambda word: (
            int(word.get("block_no", 0)),
            int(word.get("line_no", 0)),
            int(word.get("word_no", 0)),
            float(word["y0"]),
            float(word["x0"]),
        )
    )
    return words


def _text_similarity(left: str, right: str) -> float:
    a = _normalize_text(left).casefold()
    b = _normalize_text(right).casefold()

    if not a or not b:
        return 0.0

    if a == b:
        return 1.0

    if len(a) >= 20 and (a in b or b in a):
        shorter = min(len(a), len(b))
        longer = max(len(a), len(b))
        return 0.86 + (0.14 * shorter / longer)

    return SequenceMatcher(None, a, b).ratio()


def map_review_paragraphs_to_preview(
    source_path: str | Path,
    preview_pdf: str | Path,
    paragraphs: list[dict] | None = None,
) -> list[dict]:
    source = resolve_storage_path(source_path)
    preview = Path(preview_pdf)
    rows = paragraphs or extract_review_paragraphs(source)
    blocks = _pdf_text_blocks(preview)

    if source.suffix.lower() == ".pdf":
        anchors: list[dict] = []
        for row, block in zip(rows, blocks):
            anchors.append(
                {
                    "paragraph_index": int(row["paragraph_index"]),
                    "page_index": int(block["page_index"]),
                    "bbox": block["bbox"],
                    "words": block.get("words") or [],
                    "text": row["text"],
                    "section": row.get("section") or "",
                }
            )
        return anchors

    anchors: list[dict] = []
    search_start = 0

    for row in rows:
        text_value = row["text"]
        if len(text_value) < 2:
            continue

        best_start: int | None = None
        best_end: int | None = None
        best_score = 0.0

        window_end = min(len(blocks), search_start + 45)

        for start_index in range(search_start, window_end):
            start_page = int(blocks[start_index]["page_index"])
            combined_parts: list[str] = []

            for end_index in range(
                start_index,
                min(window_end, start_index + 8),
            ):
                if int(blocks[end_index]["page_index"]) != start_page:
                    break

                combined_parts.append(blocks[end_index]["text"])
                combined_text = " ".join(combined_parts)
                score = _text_similarity(text_value, combined_text)

                length_ratio = (
                    min(len(_normalize_text(text_value)), len(_normalize_text(combined_text)))
                    / max(len(_normalize_text(text_value)), len(_normalize_text(combined_text)))
                )
                adjusted_score = score * (0.82 + 0.18 * length_ratio)

                if adjusted_score > best_score:
                    best_score = adjusted_score
                    best_start = start_index
                    best_end = end_index

                if score >= 0.985 and length_ratio >= 0.9:
                    break

        if (
            best_start is None
            or best_end is None
            or best_score < 0.46
        ):
            continue

        matched_blocks = blocks[best_start : best_end + 1]
        anchor = {
            "paragraph_index": int(row["paragraph_index"]),
            "page_index": int(matched_blocks[0]["page_index"]),
            "bbox": _union_bbox(matched_blocks),
            "words": _merge_words(matched_blocks),
            "text": text_value,
            "section": row.get("section") or "",
        }
        anchors.append(anchor)

        search_start = max(search_start, best_end + 1)

    return anchors

