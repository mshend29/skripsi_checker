from __future__ import annotations

import base64
import hashlib
import html
import json
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


BASE_DIR = Path(__file__).resolve().parents[2]
HTML_PREVIEW_DIR = BASE_DIR / "storage" / "html_previews"

A4_WIDTH_MM = 210
A4_HEIGHT_MM = 297
PAGE_MARGIN_TOP_MM = 24
PAGE_MARGIN_RIGHT_MM = 24
PAGE_MARGIN_BOTTOM_MM = 24
PAGE_MARGIN_LEFT_MM = 28


def _css_length(value) -> str | None:
    if value is None:
        return None
    try:
        return f"{value.pt:.2f}pt"
    except Exception:
        return None


def _escape_text(value: str) -> str:
    return (
        html.escape(value or "")
        .replace("\t", '<span class="doc-tab">    </span>')
        .replace("\n", "<br>")
    )


def _alignment_css(paragraph) -> str:
    alignment = paragraph.alignment
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }
    return mapping.get(alignment, "left")


def _paragraph_css(paragraph) -> str:
    fmt = paragraph.paragraph_format
    styles = [
        f"text-align:{_alignment_css(paragraph)}",
        "position:relative",
    ]

    for css_name, value in (
        ("margin-left", fmt.left_indent),
        ("margin-right", fmt.right_indent),
        ("text-indent", fmt.first_line_indent),
        ("margin-top", fmt.space_before),
        ("margin-bottom", fmt.space_after),
    ):
        css_value = _css_length(value)
        if css_value:
            styles.append(f"{css_name}:{css_value}")

    line_spacing = fmt.line_spacing
    if isinstance(line_spacing, float):
        styles.append(f"line-height:{line_spacing:.3f}")
    else:
        css_value = _css_length(line_spacing)
        if css_value:
            styles.append(f"line-height:{css_value}")

    return ";".join(styles)


def _run_css(run, paragraph) -> str:
    styles: list[str] = []
    font = run.font

    if run.bold:
        styles.append("font-weight:700")
    if run.italic:
        styles.append("font-style:italic")
    if run.underline:
        styles.append("text-decoration:underline")

    size = _css_length(font.size)
    if size:
        styles.append(f"font-size:{size}")

    font_name = font.name
    if not font_name and paragraph.style is not None:
        font_name = paragraph.style.font.name
    if font_name:
        escaped = font_name.replace('"', "")
        styles.append(f'font-family:"{escaped}"')

    if font.color is not None and font.color.rgb is not None:
        styles.append(f"color:#{font.color.rgb}")

    if font.superscript:
        styles.append("vertical-align:super")
        styles.append("font-size:0.8em")
    elif font.subscript:
        styles.append("vertical-align:sub")
        styles.append("font-size:0.8em")

    if font.all_caps:
        styles.append("text-transform:uppercase")
    if font.small_caps:
        styles.append("font-variant:small-caps")

    return ";".join(styles)


def _run_images(document, run) -> list[str]:
    images: list[str] = []

    for node in run._r.iter():
        if node.tag != qn("a:blip"):
            continue

        rel_id = node.get(qn("r:embed"))
        if not rel_id:
            continue

        part = run.part.related_parts.get(rel_id)
        if part is None:
            continue

        mime = getattr(part, "content_type", "image/png")
        encoded = base64.b64encode(part.blob).decode("ascii")

        width_css = ""
        for drawing_node in run._r.iter():
            if drawing_node.tag == qn("wp:extent"):
                try:
                    width_px = int(int(drawing_node.get("cx")) / 9525)
                    if width_px > 0:
                        width_css = f"width:{width_px}px;"
                except Exception:
                    pass
                break

        images.append(
            '<img class="doc-inline-image" '
            f'src="data:{mime};base64,{encoded}" '
            f'style="{width_css}max-width:100%;height:auto;" alt="">'
        )

    return images


def _iter_paragraph_runs(paragraph):
    if hasattr(paragraph, "iter_inner_content"):
        for item in paragraph.iter_inner_content():
            if hasattr(item, "_r"):
                yield item
                continue

            for run in getattr(item, "runs", []):
                yield run
        return

    yield from paragraph.runs


def _paragraph_has_page_break(paragraph) -> bool:
    for run in _iter_paragraph_runs(paragraph):
        for node in run._r.iter():
            if node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
                return True
    return False


def _render_runs(document, paragraph) -> str:
    parts: list[str] = []

    for run in _iter_paragraph_runs(paragraph):
        css = _run_css(run, paragraph)
        text_value = _escape_text(run.text)
        if text_value:
            parts.append(
                f'<span class="doc-run" style="{css}">{text_value}</span>'
            )
        parts.extend(_run_images(document, run))

    return "".join(parts) if parts else "&nbsp;"


def _render_paragraph(
    document,
    paragraph,
    paragraph_index: int | None,
    reviewable: bool = True,
) -> str:
    css = _paragraph_css(paragraph)
    style_name = paragraph.style.name if paragraph.style is not None else ""

    classes = ["flow-block", "doc-paragraph"]
    if style_name.lower().startswith("heading"):
        classes.append("doc-heading")

    attrs: list[str] = []
    if paragraph_index is not None and reviewable:
        classes.append("review-paragraph")
        attrs.extend(
            [
                f'data-paragraph-index="{paragraph_index}"',
                f'id="paragraph-{paragraph_index}"',
            ]
        )

    if paragraph.paragraph_format.page_break_before:
        attrs.append('data-page-break-before="1"')
    if _paragraph_has_page_break(paragraph):
        attrs.append('data-page-break-after="1"')

    content = _render_runs(document, paragraph)

    return (
        f'<div class="{" ".join(classes)}" {" ".join(attrs)} '
        f'style="{css}">'
        '<div class="paragraph-content">'
        f"{content}"
        "</div></div>"
    )


def _cell_colspan(cell) -> int:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1

    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return 1

    try:
        return max(1, int(grid_span.get(qn("w:val"))))
    except Exception:
        return 1


def _render_table(document, table) -> str:
    rows_html: list[str] = []

    for row in table.rows:
        cells_html: list[str] = []
        seen_cells: set[int] = set()

        for cell in row.cells:
            identity = id(cell._tc)
            if identity in seen_cells:
                continue
            seen_cells.add(identity)

            colspan = _cell_colspan(cell)
            body = "".join(
                _render_paragraph(
                    document,
                    paragraph,
                    paragraph_index=None,
                    reviewable=False,
                )
                for paragraph in cell.paragraphs
            )

            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ""
            cells_html.append(f"<td{colspan_attr}>{body}</td>")

        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")

    return (
        '<div class="flow-block table-block">'
        '<table class="doc-table">'
        + "".join(rows_html)
        + "</table></div>"
    )


def _header_footer_html(document) -> tuple[str, str]:
    if not document.sections:
        return "", ""

    section = document.sections[0]

    header = "".join(
        _render_paragraph(
            document,
            paragraph,
            paragraph_index=None,
            reviewable=False,
        )
        for paragraph in section.header.paragraphs
        if paragraph.text.strip()
    )

    footer = "".join(
        _render_paragraph(
            document,
            paragraph,
            paragraph_index=None,
            reviewable=False,
        )
        for paragraph in section.footer.paragraphs
        if paragraph.text.strip()
    )

    return header, footer


def _body_blocks(document) -> list[str]:
    paragraph_iter = iter(enumerate(document.paragraphs))
    table_iter = iter(document.tables)
    blocks: list[str] = []

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            try:
                paragraph_index, paragraph = next(paragraph_iter)
            except StopIteration:
                continue

            blocks.append(
                _render_paragraph(
                    document,
                    paragraph,
                    paragraph_index,
                    reviewable=True,
                )
            )

        elif child.tag == qn("w:tbl"):
            try:
                table = next(table_iter)
            except StopIteration:
                continue

            blocks.append(_render_table(document, table))

    return blocks


def _page_metrics(document) -> dict[str, float]:
    if not document.sections:
        return {
            "width": A4_WIDTH_MM,
            "height": A4_HEIGHT_MM,
            "top": PAGE_MARGIN_TOP_MM,
            "right": PAGE_MARGIN_RIGHT_MM,
            "bottom": PAGE_MARGIN_BOTTOM_MM,
            "left": PAGE_MARGIN_LEFT_MM,
        }

    section = document.sections[0]

    def mm(value, fallback):
        try:
            return float(value.mm)
        except Exception:
            return float(fallback)

    return {
        "width": mm(section.page_width, A4_WIDTH_MM),
        "height": mm(section.page_height, A4_HEIGHT_MM),
        "top": mm(section.top_margin, PAGE_MARGIN_TOP_MM),
        "right": mm(section.right_margin, PAGE_MARGIN_RIGHT_MM),
        "bottom": mm(section.bottom_margin, PAGE_MARGIN_BOTTOM_MM),
        "left": mm(section.left_margin, PAGE_MARGIN_LEFT_MM),
    }


def render_docx_html(
    file_path: str | Path,
    comment_counts: dict[int, int] | None = None,
) -> str:
    path = Path(file_path)
    if path.suffix.lower() != ".docx":
        raise ValueError("Renderer HTML interaktif membutuhkan file DOCX.")

    document = DocxDocument(path)
    page = _page_metrics(document)
    header_html, footer_html = _header_footer_html(document)
    blocks = _body_blocks(document)
    counts_json = json.dumps(comment_counts or {})

    normal_style = (
        document.styles["Normal"]
        if "Normal" in document.styles
        else None
    )
    default_font_name = (
        normal_style.font.name
        if normal_style is not None and normal_style.font.name
        else "Times New Roman"
    )
    default_font_size = (
        normal_style.font.size.pt
        if normal_style is not None and normal_style.font.size is not None
        else 12
    )

    css = f"""
    html, body {{
        margin: 0;
        padding: 0;
        background: #d8d8d8;
        color: #000;
        font-family: "{html.escape(default_font_name)}", "Times New Roman", serif;
        font-size: {default_font_size:.1f}pt;
    }}
    * {{ box-sizing: border-box; }}
    body {{ overflow-x: auto; }}
    #pages {{
        padding: 24px 28px 48px;
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 22px;
    }}
    #source-flow {{
        position: absolute;
        left: -100000px;
        top: 0;
        width: {page["width"] - page["left"] - page["right"]:.2f}mm;
        visibility: hidden;
    }}
    .word-page {{
        width: {page["width"]:.2f}mm;
        height: {page["height"]:.2f}mm;
        background: #fff;
        box-shadow: 0 2px 10px rgba(0,0,0,.22);
        position: relative;
        flex: none;
        overflow: hidden;
    }}
    .page-header {{
        position: absolute;
        left: {page["left"]:.2f}mm;
        right: {page["right"]:.2f}mm;
        top: 8mm;
        min-height: 10mm;
        font-size: 9pt;
    }}
    .page-body {{
        position: absolute;
        left: {PAGE_MARGIN_LEFT_MM}mm;
        right: {PAGE_MARGIN_RIGHT_MM}mm;
        top: {page["top"]:.2f}mm;
        bottom: {page["bottom"]:.2f}mm;
        overflow: hidden;
    }}
    .page-footer {{
        position: absolute;
        left: {PAGE_MARGIN_LEFT_MM}mm;
        right: {PAGE_MARGIN_RIGHT_MM}mm;
        bottom: 7mm;
        min-height: 9mm;
        font-size: 9pt;
    }}
    .flow-block {{ width: 100%; }}
    .doc-paragraph {{
        min-height: 1.15em;
        padding: 2px 4px;
        margin-left: -4px;
        margin-right: -4px;
        border: 1px solid transparent;
        border-radius: 4px;
        transition: background .12s ease, border-color .12s ease;
    }}
    .doc-paragraph .paragraph-content {{ min-height: 1em; }}
    .review-paragraph {{ cursor: text; }}
    .review-paragraph:hover {{
        border-color: #b8b8b8;
        background: rgba(247,247,247,.7);
    }}
    .review-paragraph.review-active {{
        border-color: #777;
        background: rgba(255,248,195,.45);
    }}
    .comment-action {{
        position: absolute;
        z-index: 20;
        display: none;
        background: #fff;
        color: #111;
        border: 1px solid #aaa;
        border-radius: 5px;
        padding: 5px 9px;
        font-family: Arial, sans-serif;
        font-size: 11px;
        cursor: pointer;
        box-shadow: 0 2px 6px rgba(0,0,0,.12);
        white-space: nowrap;
    }}
    .review-paragraph:hover > .comment-action,
    .review-paragraph.review-active > .comment-action {{
        display: block;
        right: 4px;
        top: 3px;
    }}
    .comment-badge {{
        position: absolute;
        right: -23px;
        top: 3px;
        min-width: 19px;
        height: 19px;
        padding: 0 5px;
        border-radius: 10px;
        background: #fff;
        border: 1px solid #aaa;
        color: #111;
        font: 10px/17px Arial, sans-serif;
        text-align: center;
        z-index: 19;
    }}
    .doc-run {{ white-space: pre-wrap; }}
    .doc-inline-image {{
        display: inline-block;
        vertical-align: baseline;
    }}
    .table-block {{ margin: 8px 0 10px; }}
    .doc-table {{
        border-collapse: collapse;
        width: 100%;
        table-layout: auto;
        color: #000;
        background: #fff;
    }}
    .doc-table td, .doc-table th {{
        border: 1px solid #444;
        padding: 5px 7px;
        vertical-align: top;
    }}
    .doc-table .doc-paragraph {{
        padding: 0;
        margin: 0;
        border: 0;
    }}
    ::selection {{
        background: #89bfff;
        color: #000;
    }}
    """

    js = f"""
    const COMMENT_COUNTS = {counts_json};
    let lastSelection = null;

    function newPage() {{
        const page = document.createElement('section');
        page.className = 'word-page';

        const header = document.createElement('div');
        header.className = 'page-header';
        header.innerHTML = document.getElementById('header-template').innerHTML;

        const body = document.createElement('div');
        body.className = 'page-body';

        const footer = document.createElement('div');
        footer.className = 'page-footer';
        footer.innerHTML = document.getElementById('footer-template').innerHTML;

        page.appendChild(header);
        page.appendChild(body);
        page.appendChild(footer);
        document.getElementById('pages').appendChild(page);
        return body;
    }}

    function paginate() {{
        const source = document.getElementById('source-flow');
        const blocks = Array.from(source.children);
        let pageBody = newPage();

        for (const block of blocks) {{
            const forceBefore = block.dataset.pageBreakBefore === '1';

            if (forceBefore && pageBody.children.length) {{
                pageBody = newPage();
            }}

            pageBody.appendChild(block);

            if (
                pageBody.scrollHeight > pageBody.clientHeight
                && pageBody.children.length > 1
            ) {{
                pageBody.removeChild(block);
                pageBody = newPage();
                pageBody.appendChild(block);
            }}

            if (block.dataset.pageBreakAfter === '1') {{
                pageBody = newPage();
            }}
        }}

        source.remove();

        const pages = document.querySelectorAll('.word-page');
        if (
            pages.length > 1
            && pages[pages.length - 1].querySelector('.page-body').children.length === 0
        ) {{
            pages[pages.length - 1].remove();
        }}
    }}

    function paragraphFromNode(node) {{
        if (!node) return null;
        const element = node.nodeType === Node.ELEMENT_NODE
            ? node
            : node.parentElement;
        return element ? element.closest('.review-paragraph') : null;
    }}

    function updateSelection() {{
        const selection = window.getSelection();
        if (!selection || selection.rangeCount === 0 || selection.isCollapsed) {{
            return;
        }}

        const text = selection.toString().trim();
        if (!text) return;

        const startParagraph = paragraphFromNode(selection.anchorNode);
        const endParagraph = paragraphFromNode(selection.focusNode);

        if (
            !startParagraph
            || !endParagraph
            || startParagraph !== endParagraph
        ) {{
            return;
        }}

        lastSelection = {{
            paragraphIndex: Number(startParagraph.dataset.paragraphIndex),
            text: text
        }};
    }}

    function bindParagraph(paragraph) {{
        const index = Number(paragraph.dataset.paragraphIndex);

        const button = document.createElement('button');
        button.className = 'comment-action';
        button.type = 'button';
        button.textContent = '+ Tambahkan komentar';

        button.addEventListener('mousedown', (event) => {{
            event.preventDefault();
        }});

        button.addEventListener('click', (event) => {{
            event.preventDefault();
            event.stopPropagation();

            const selected = (
                lastSelection
                && lastSelection.paragraphIndex === index
                && lastSelection.text
            )
                ? lastSelection.text
                : paragraph.querySelector('.paragraph-content').innerText.trim();

            if (window.reviewBridge) {{
                window.reviewBridge.requestComment(index, selected);
            }}
        }});

        paragraph.appendChild(button);

        const count = Number(COMMENT_COUNTS[String(index)] || 0);
        if (count > 0) {{
            const badge = document.createElement('span');
            badge.className = 'comment-badge';
            badge.textContent = String(count);
            paragraph.appendChild(badge);
        }}
    }}

    function setupReview() {{
        document.querySelectorAll('.review-paragraph').forEach(bindParagraph);

        document.addEventListener('mouseup', () => {{
            window.setTimeout(updateSelection, 0);
        }});

        document.addEventListener('keyup', () => {{
            window.setTimeout(updateSelection, 0);
        }});
    }}

    function scrollToParagraph(index) {{
        document.querySelectorAll('.review-active').forEach((node) => {{
            node.classList.remove('review-active');
        }});

        const node = document.getElementById('paragraph-' + String(index));
        if (!node) return false;

        node.classList.add('review-active');
        node.scrollIntoView({{
            behavior: 'smooth',
            block: 'center'
        }});
        return true;
    }}

    window.scrollToParagraph = scrollToParagraph;

    window.addEventListener('DOMContentLoaded', () => {{
        new QWebChannel(qt.webChannelTransport, (channel) => {{
            window.reviewBridge = channel.objects.reviewBridge;
        }});

        const imageWaiters = Array.from(document.images).map((image) => {{
            if (image.complete) return Promise.resolve();

            return new Promise((resolve) => {{
                image.addEventListener('load', resolve, {{ once: true }});
                image.addEventListener('error', resolve, {{ once: true }});
            }});
        }});

        const fontsReady = document.fonts
            ? document.fonts.ready
            : Promise.resolve();

        Promise.all([fontsReady, ...imageWaiters]).then(() => {{
            paginate();
            setupReview();
        }});
    }});
    """

    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<meta name="color-scheme" content="light">
<style>{css}</style>
<script src="qrc:///qtwebchannel/qwebchannel.js"></script>
</head>
<body>
<div id="header-template" hidden>{header_html}</div>
<div id="footer-template" hidden>{footer_html}</div>
<div id="source-flow">{"".join(blocks)}</div>
<div id="pages"></div>
<script>{js}</script>
</body>
</html>
"""



def write_docx_html_preview(
    file_path: str | Path,
    comment_counts: dict[int, int] | None = None,
) -> Path:
    source = Path(file_path)
    stat = source.stat()
    payload = (
        f"{source.resolve()}|{stat.st_size}|{stat.st_mtime_ns}|"
        + json.dumps(comment_counts or {}, sort_keys=True)
    )
    digest = hashlib.sha1(payload.encode("utf-8")).hexdigest()[:16]

    HTML_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    target = HTML_PREVIEW_DIR / f"{source.stem}_{digest}.html"

    if not target.exists():
        target.write_text(
            render_docx_html(
                source,
                comment_counts=comment_counts,
            ),
            encoding="utf-8",
        )

    return target
